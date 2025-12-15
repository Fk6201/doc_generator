from docx import Document
from docx.shared import Pt
import os
from typing import List
from models import ReportItem

class ReportGenerator:
    def __init__(self, output_dir: str = "output_reports"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate(self, filename: str, title: str, items: List[ReportItem]) -> str:
        doc = Document()
        doc.add_heading(title, level=1)

        for idx, item in enumerate(items, 1):
            doc.add_heading(f"{idx}. {item.title}", level=2)
            doc.add_paragraph(item.summary)
            if item.value is not None:
                p = doc.add_paragraph()
                run = p.add_run(f"Value: {item.value}")
                run.bold = True
                run.font.size = Pt(11)

        out_path = os.path.join(self.output_dir, filename)
        doc.save(out_path)
        return out_path
